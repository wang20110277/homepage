import json
from docx import Document
import os


def load_json(file_path):
    """加载 JSON 文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    except Exception as e:
        print(f"加载 JSON 文件失败: {e}")
        return {}


def extract_fields(json_data):
    """从疑似实际控制人数据中提取相关字段"""
    extracted_data = {}

    # 提取 actualControllerList
    actual_controller_list = (json_data.get('result') or {}).get('actualControllerList', [])
    extracted_data['actualControllerList'] = [
        {
            'hId': item.get('hId', ''),
            'hPid': item.get('hPid', ''),
            'gId': item.get('gId', ''),
            'name': item.get('name', ''),
            'type': item.get('type', ''),
            'ratio': item.get('ratio', '')
        }
        for item in actual_controller_list
    ]

    # 提取 relationships
    relationships = []
    for path_key, path_value in (json_data.get('result') or {}).get('pathMap', {}).items():
        for relationship in path_value.get('relationships', []):
            relationships.append({
                'startNode': relationship.get('startNode', ''),
                'type': relationship.get('type', ''),
                'endNode': relationship.get('endNode', ''),
                'percentStr': relationship.get('properties', {}).get('percentStr', ''),
                'isRed': relationship.get('properties', {}).get('isRed', ''),
                'percent': relationship.get('properties', {}).get('percent', '')
            })
    extracted_data['relationships'] = relationships

    # 提取 nodes
    nodes = []
    for path_key, path_value in (json_data.get('result') or {}).get('pathMap', {}).items():
        for node in path_value.get('nodes', []):
            node_info = {
                'id': node.get('id', ''),
                'labels': node.get('labels', [])
            }
            # 当labels为Human时，提取额外字段
            if 'Human' in node.get('labels', []):
                node_info['name'] = node.get('properties', {}).get('name', '')
                node_info['type'] = node.get('properties', {}).get('type', '')
                node_info['cid'] = node.get('properties', {}).get('cid', '')
            elif 'Company' in node.get('labels', []):
                node_info['name'] = node.get('properties', {}).get('name', '')
                node_info['type'] = node.get('properties', {}).get('type', '')
                node_info['gid'] = node.get('properties', {}).get('gid', '')
            nodes.append(node_info)
    extracted_data['nodes'] = nodes

    print("提取的实际控制人字段:")
    print(json.dumps(extracted_data, indent=2, ensure_ascii=False))
    return extracted_data


def replace_placeholders(doc, data):
    """替换文档中的占位符"""
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"【{key}】"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))
    return doc


def add_text_after_specific_text(doc, search_text, new_text):
    """在指定文本后添加新文本"""
    found = False
    for para in doc.paragraphs:
        if search_text in para.text and not found:
            found = True
            para.add_run('\n' + new_text)
    return found


def replace_text_in_table(doc, old_text_list, new_text_list):
    """替换表格中的文本"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for i in range(len(old_text_list)):
                        # 确保替换值是字符串
                        old_text = old_text_list[i]
                        new_text = str(
                            new_text_list[i]) if new_text_list[i] is not None else ""
                        paragraph.text = paragraph.text.replace(
                            old_text, new_text)


def create_report(output_path, word_path_file):
    # 加载疑似实际控制人数据
    control_json = output_path + './疑似实际控制人.json'
    control_data = load_json(control_json)
    extracted_fields = extract_fields(control_data)

    # 加载主数据
    with open(output_path + '/extracted_data.json', 'r', encoding='utf-8') as f:
        main_data = json.load(f)

    # 加载Word文档
    # 使用相对路径获取template.docx
    current_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(current_dir, 'template.docx')
    if not os.path.exists(doc_path):
        print(f"错误: Word文档不存在: {doc_path}")
        return
    doc = Document(doc_path)

    # 准备替换数据
    business_info = main_data["工商信息"]["result"]

    # 安全获取股东和评级信息
    shareholders = main_data.get("企业股东", [])
    credit_ratings = main_data.get("企业信用评级", [])

    # 安全获取实际控制人数据
    actual_controllers = extracted_fields.get('actualControllerList', [])
    relationships = extracted_fields.get('relationships', [])
    nodes = extracted_fields.get('nodes', [])

    # 获取第一条实际控制人数据（如果存在）
    first_controller = actual_controllers[0] if actual_controllers else {}
    first_relationship = relationships[0] if relationships else {}

    # 从nodes中找到第一个Company类型的节点
    first_company_node = {}
    for node in nodes:
        if 'Company' in node.get('labels', []):
            first_company_node = node
            break

    replacements = {
        # 工商信息字段
        "name": business_info.get("name", "") if business_info.get("name") != 'null' else "",
        "establishTime": business_info.get("estabishTime", "") if business_info.get("estabishTime") != 'null' else "",
        "creditCode": business_info.get("creditCode", "") if business_info.get("creditCode") != 'null' else "",
        "orgNumber": business_info.get("orgNumber", "") if business_info.get("orgNumber") != 'null' else "",
        "companyOrgType": business_info.get("companyOrgType", "") if business_info.get("companyOrgType") != 'null' else "",
        "regInstitute": business_info.get("regInstitute", "") if business_info.get("regInstitute") != 'null' else "",
        "approvedTime": business_info.get("approvedTime", "") if business_info.get("approvedTime") != 'null' else "",
        "regLocation": business_info.get("regLocation", "") if business_info.get("regLocation") != 'null' else "",
        "regCapital": business_info.get("regCapital", "") if business_info.get("regCapital") != 'null' else "",
        "actualCapital": business_info.get("actualCapital", "") if business_info.get("actualCapital") != 'null' else "",
        "totalnumber": shareholders[0].get('totalNumber', "") if shareholders else "",
        "total": credit_ratings[0].get('total', "") if credit_ratings else "",

        # 实际控制人字段 (actualControllerList)
        "hId": str(first_controller.get('hId', "")) if first_controller.get('hId') else "",
        "hPid": str(first_controller.get('hPid', "")) if first_controller.get('hPid') else "",
        "gId": str(first_controller.get('gId', "")) if first_controller.get('gId') else "",
        "type": str(first_controller.get('type', "")) if first_controller.get('type') else "",
        "ratio": str(first_controller.get('ratio', "")) if first_controller.get('ratio') else "",

        # 关系字段 (relationships)
        "startNode": str(first_relationship.get('startNode', "")) if first_relationship.get('startNode') else "",
        "endNode": str(first_relationship.get('endNode', "")) if first_relationship.get('endNode') else "",
        "percent": str(first_relationship.get('percent', "")) if first_relationship.get('percent') else "",
        "isRed": str(first_relationship.get('isRed', "")) if first_relationship.get('isRed') != '' else "",
        "percentStr": str(first_relationship.get('percentStr', "")) if first_relationship.get('percentStr') else "",

        # 节点字段 (nodes - Company类型)
        "gid": str(first_company_node.get('gid', "")) if first_company_node.get('gid') else "",
        "labels": ', '.join(first_company_node.get('labels', [])) if first_company_node.get('labels') else "",
        "nodeType": str(first_company_node.get('type', "")) if first_company_node.get('type') else ""
    }

    # 执行段落替换
    doc = replace_placeholders(doc, replacements)

    # 添加主要成员信息
    staff_list = main_data.get('工商信息', {}).get('staffList', [])
    if staff_list:
        for line in staff_list:
            type_join = line.get('typeJoin', [''])
            # 处理typeJoin为'null'字符串的情况
            if type_join == 'null' or not type_join:
                type_join = ['']
            i_str = f"{type_join[0]}：{line.get('name', '')};{line.get('id', '')};"
            if not add_text_after_specific_text(doc, '2、主要成员', i_str):
                print(f"警告: 未找到 '2、主要成员' 文本")
    else:
        print("提示: 工商信息中无主要成员数据")

    # 添加股东信息
    for line in main_data.get("企业股东", []):
        capital = line.get("capital", [{}])[0] if line.get("capital") else {}
        capital_actl = line.get("capitalActl", [{}])[0] if line.get("capitalActl") else {}
        i_str = (f"出资人：{line.get('name', '')}；"
                 f"简称：{line.get('alias', '')}；"
                 f"类型：{line.get('type', '')}；"
                 f"股份占比：{capital.get('percent', '')}；"
                 f"出资金额：{capital.get('amomon', '')}；"
                 f"出资时间：{capital.get('time', '')}；"
                 f"认缴方式：{capital.get('payment', '')}；"
                 f"实缴：{capital_actl.get('amomon', '')}")
        if not add_text_after_specific_text(doc, '股东总数', i_str):
            print(f"警告: 未找到 '股东总数' 文本")

    # 添加信用评级信息
    credit_ratings = main_data.get('企业信用评级', [])
    if credit_ratings:
        for line in credit_ratings:
            # 安全获取各字段，处理null值
            rating_company = line.get('ratingCompanyName', '') if line.get('ratingCompanyName') != 'null' else ''
            alias = line.get('alias', '') if line.get('alias') != 'null' else ''
            gid = line.get('gid', '') if line.get('gid') != 'null' else ''
            rating_date = line.get('ratingDate', '') if line.get('ratingDate') != 'null' else ''
            rating_outlook = line.get('ratingOutlook', '') if line.get('ratingOutlook') != 'null' else ''
            subject_level = line.get('subjectLevel', '') if line.get('subjectLevel') != 'null' else ''

            i_str = (f"评级公司{rating_company},"
                     f"简称{alias},"
                     f"公司id{gid}。"
                     f"在{rating_date}出具评级结果，"
                     f"评级展望为{rating_outlook},"
                     f"评级等级为{subject_level}")
            if not add_text_after_specific_text(doc, '经天眼查网站查询', i_str):
                print(f"警告: 未找到 '经天眼查网站查询' 文本")
    else:
        print("提示: 无企业信用评级数据")

    # 准备表格替换数据
    table_old_text = [
        '【result对象name】', '【legalPersonName】', '【regCapital】', '【creditCode】',
        '【establishTime】', '【taxNumber】', '【companyOrgType】', '【orgNumber】',
        '【approvedTime】', '【industry】', '【base】', '【regInstitute】',
        '【fromTime】', '【revokeDate】', '【regLocation】', '【businessScope】'
    ]

    table_new_text = [
        business_info.get('name', '') if business_info.get('name') != 'null' else '',
        business_info.get('legalPersonName', '') if business_info.get('legalPersonName') != 'null' else '',
        business_info.get('regCapital', '') if business_info.get('regCapital') != 'null' else '',
        business_info.get('creditCode', '') if business_info.get('creditCode') != 'null' else '',
        business_info.get('estabishTime', '') if business_info.get('estabishTime') != 'null' else '',
        business_info.get('taxNumber', '') if business_info.get('taxNumber') != 'null' else '',
        business_info.get('companyOrgType', '') if business_info.get('companyOrgType') != 'null' else '',
        business_info.get('orgNumber', '') if business_info.get('orgNumber') != 'null' else '',
        business_info.get('approvedTime', '') if business_info.get('approvedTime') != 'null' else '',
        business_info.get('industry', '') if business_info.get('industry') != 'null' else '',
        business_info.get('base', '') if business_info.get('base') != 'null' else '',
        business_info.get('regInstitute', '') if business_info.get('regInstitute') != 'null' else '',
        business_info.get('fromTime', '') if business_info.get('fromTime') != 'null' else '',
        business_info.get('revokeDate', '') if business_info.get('revokeDate') != 'null' else '',
        business_info.get('regLocation', '') if business_info.get('regLocation') != 'null' else '',
        business_info.get('businessScope', '') if business_info.get('businessScope') != 'null' else ''
    ]

    # 执行表格文本替换
    replace_text_in_table(doc, table_old_text, table_new_text)

    # 处理表格数据 - 使用安全访问
    tables = doc.tables
    print(f"文档中包含 {len(tables)} 个表格")

    # 表格1: 司法风险
    # 修改表格1: 司法风险部分的代码
    if len(tables) > 1:
        table1 = tables[1]
        law_suits = main_data.get("司法风险", {}).get("lawSuitList", [])
        kt_announcements = main_data.get(
            "司法风险", {}).get("ktAnnouncementList", [])

        # 确保列表不为null或空
        if law_suits and kt_announcements:
            min_length = min(len(law_suits), len(kt_announcements))
            for i in range(min_length):
                new_row = table1.add_row()
                # 安全访问被告信息
                defendant = law_suits[i].get("defendants", "未知被告")
                # 处理被告信息为null的情况
                if defendant == 'null' or defendant is None:
                    defendant = "未知被告"

                # 安全访问原告信息 - 修改后的部分
                plaintiffs = kt_announcements[i].get("plaintiffs", [])
                plaintiff_name = "未知原告"
                if plaintiffs and plaintiffs != 'null':  # 确保plaintiffs列表不为空且不为'null'字符串
                    if isinstance(plaintiffs, list) and len(plaintiffs) > 0:
                        plaintiff_name = plaintiffs[0].get("name", "未知原告")

                new_row.cells[0].text = f'原告：{plaintiff_name}\n被告：{defendant}'
                new_row.cells[1].text = law_suits[i].get("casetype", "")
                new_row.cells[2].text = law_suits[i].get("casereason", "")
        else:
            print("提示: 该企业暂无司法风险数据")
    else:
        print("警告: 文档缺少第二个表格（司法风险）")

    # 表格2: 工商变更信息
    if len(tables) > 2:
        table2 = tables[2]
        changes = main_data["工商信息"].get("changeList", [])
        table2.cell(0, 0).text = f"共涉及{len(changes)}条变更，明细如下："

        for change in changes:
            new_row = table2.add_row()
            new_row.cells[0].text = change.get('changeTime', "")
            new_row.cells[1].text = change.get('changeItem', "")
            new_row.cells[2].text = change.get('contentBefore', "")
            new_row.cells[3].text = change.get('contentAfter', "")
    else:
        print("警告: 文档缺少第三个表格（工商变更）")

    # 表格3: 控股股东
    if len(tables) > 3:
        table3 = tables[3]
        shareholders = main_data.get("企业股东", [])

        for i, shareholder in enumerate(shareholders):
            capitals = shareholder.get("capital", [])
            for capital in capitals:
                new_row = table3.add_row()
                new_row.cells[0].text = str(i + 1)
                new_row.cells[1].text = shareholder.get("name", "")
                new_row.cells[2].text = capital.get("percent", "")
                new_row.cells[3].text = capital.get("amomon", "")
                new_row.cells[4].text = capital.get("time", "")
    else:
        print("警告: 文档缺少第四个表格（控股股东）")

    # 表格4: 实际控制人
    if len(tables) > 4:
        table4 = tables[4]
        controllers = extracted_fields.get("actualControllerList", [])

        for i, controller in enumerate(controllers):
            new_row = table4.add_row()
            new_row.cells[0].text = str(controller.get('type', ''))
            new_row.cells[1].text = controller.get('name', '')
            new_row.cells[2].text = str(controller.get('ratio', ''))
    else:
        print("警告: 文档缺少第五个表格（实际控制人）")

    # 表格5: 主要人员
    if len(tables) > 5:
        table5 = tables[5]
        staff_items = main_data.get("主要人员", {}).get("items", [])
        table5.cell(0, 0).text = f"主要人员总数：{len(staff_items)}"

        for staff in staff_items:
            new_row = table5.add_row()
            new_row.cells[0].text = staff.get('typeJoin', [''])[0]  # 取第一个职位
            new_row.cells[1].text = staff.get('name', '')
            new_row.cells[2].text = staff.get('hcgid', '')
    else:
        print("警告: 文档缺少第六个表格（主要人员）")

    # 保存文档
    doc.save(word_path_file)
    print(f"文档处理完成，已保存至: {word_path_file}")
